#!/usr/bin/env python3
"""Generate test fixtures for parser tests."""

from pathlib import Path
from docx import Document
import io

FIXTURES = Path(__file__).parent / "fixtures"
FIXTURES.mkdir(exist_ok=True)

print(f"Creating test fixtures in {FIXTURES}")

# Create sample DOCX
doc = Document()
doc.add_heading("Sample Document", 0)
doc.add_paragraph("This is a sample DOCX file for testing.")
doc.add_paragraph("It has multiple paragraphs.")
doc.add_paragraph("The parser should extract all text content.")
doc.save(FIXTURES / "sample.docx")
print("✓ Created sample.docx")

# Create sample Markdown
markdown_content = """# Sample Markdown

This is a **sample** markdown document.

## Section Two

- List item 1
- List item 2

Regular paragraph text for testing the markdown parser.
"""
(FIXTURES / "sample.md").write_text(markdown_content)
print("✓ Created sample.md")

# Create a simple PDF using reportlab if available, otherwise use pypdf
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    pdf_path = str(FIXTURES / "sample.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.drawString(100, 750, "Sample PDF Document")
    c.drawString(100, 730, "This is test content for parsing.")
    c.drawString(100, 710, "The PDF parser should extract all text.")
    c.save()
    print("✓ Created sample.pdf (using reportlab)")

except ImportError:
    print("⚠ reportlab not available, creating PDF using alternative method")
    # Create a minimal valid PDF manually
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
/Contents 4 0 R
/Resources <<
/Font <<
/F1 <<
/Type /Font
/Subtype /Type1
/BaseFont /Helvetica
>>
>>
>>
>>
endobj
4 0 obj
<<
/Length 120
>>
stream
BT
/F1 12 Tf
100 750 Td
(Sample PDF Document) Tj
0 -20 Td
(This is test content for parsing.) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000317 00000 n
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
485
%%EOF
"""
    (FIXTURES / "sample.pdf").write_bytes(pdf_content)
    print("✓ Created sample.pdf (minimal PDF)")

# Create empty PDF for testing edge case
empty_pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj
4 0 obj
<<
/Length 0
>>
stream
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000264 00000 n
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
314
%%EOF
"""
(FIXTURES / "empty.pdf").write_bytes(empty_pdf_content)
print("✓ Created empty.pdf")

print("\nAll fixtures created successfully!")
print(f"Location: {FIXTURES.absolute()}")
